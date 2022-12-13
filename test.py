from openpyxl import Workbook, load_workbook
from docx import Document
from glob import glob
import time

def get_doc_tables_dict() -> dict:
    files_amount = len(glob(r'resources\docs\*.docx'))

    for doc_num in range(files_amount):
        doc = Document(f'resources\docs\{doc_num+1}.docx')
        tables = doc.tables

        if len(tables) >= 6:
            tables = tables[:-1]

        for table_num in range(len(tables)):
            table = tables[table_num]
            date = doc.paragraphs[paragraphs[table_num]].text.split()[0]
            doc_tables[date] = table

def docx_to_excel():
    wb = Workbook()
    ws = wb.active

    ex_row = 0
    ex_col = 1

    for date in doc_tables:
        table = doc_tables[date]
        start_ex_row = ex_row
        for doc_row in table.rows:
            ex_row += 1
            for doc_cell in doc_row.cells:
                ex_col += 1
                text = doc_cell.text
                ws._get_cell(row=ex_row, column=1).value = str(date)

                if '=' not in text:
                    ws._get_cell(row=ex_row, column=ex_col).value = str(text)
                else:
                    ws._get_cell(row=ex_row, column=ex_col).value = None
            ex_col = 0
        dates[date] = [start_ex_row+1, ex_row+1]
    else: 
        wb.save(r'resources\tables\changes.xlsx')
        wb.close()

    

paragraphs = [3, 10, 18, 26, 34]
doc_tables = {}
dates = {}

st = time.time()
get_doc_tables_dict()
docx_to_excel()
print(time.time()-st)
print(dates)