from docx import Document
from openpyxl import load_workbook, Workbook
from calendar import monthrange, weekday
from teachers_list import teachers_list
from glob import glob

from data import week_days, days_strs, offset

import time
import shutil

def resave_file(src: str, dst='tabel.xlsx') -> None:
    shutil.copy2(src=src, dst=dst)

def script(gen_excel_table_path: str, document_path: str) -> None:
    start_values = get_start_day(document_path=document_path)

    days_amount = start_values[0]
    week_num = start_values[1]
    day_num = start_values[-1]
    date = start_values[2].split('.')[1:]

    resave_file(src=f'resources/tables/base{days_amount}.xlsx')

    wb_gen = load_workbook(filename=gen_excel_table_path)
    ws_gen = wb_gen.active

    wb_base = load_workbook(filename='tabel.xlsx')
    ws_base = wb_base.active

    wb_ch = load_workbook(filename=r'resources\tables\changes.xlsx')
    ws_ch = wb_ch.active

    last_full_column = 2

    for teacher in teachers_list:
        ws_base._get_cell(row=1, column=last_full_column+1).value = teacher
        last_full_column += 1
    else:
        wb_base.save(filename='tabel.xlsx')

    curr_str = 2

    for i in range(days_amount):
        ws_base[f'A{curr_str}'].value = get_day(date=f'{i+1}.{date[0]}.{date[1]}', week_num=week_num)

        if 'воскресенье' in ws_base[f'A{curr_str}'].value:
            if week_num == '1':
                week_num = '2'
            else:
                week_num = '1'

        for teacher_num in range(len(teachers_list)):
            teacher = teachers_list[teacher_num].split()[0]
            if 'суббота' not in ws_base[f'A{curr_str}'].value and 'воскресенье' not in ws_base[f'A{curr_str}'].value:
                for row in range(days_strs[week_days[day_num]], days_strs[week_days[day_num+1]]):
                    #перебираем все номера столбцов в таблице
                    for col in range(1, ws_gen.max_column+1):
                        #получение ячейки: row-№строки, col-№столбца
                        cell = ws_gen._get_cell(row=row, column=col)
                        
                        #получение значения ячейки
                        value = str(cell.value)
                        value_ed = del_all_spaces(value)

                        if teacher in value.split():
                            time = del_all_spaces(ws_gen._get_cell(row=cell.row, column=2).value)
                            group = get_group_name(get_group(del_spaces(ws_gen._get_cell(row=cell.row-1-offset[time], column=cell.column).value)))
                            curr_cell = ws_base._get_cell(row=curr_str+offset[time], column=teacher_num+3)
                            if '1н' in value_ed or '2н' in value_ed:
                                values = del_spaces(value).split('2 н')
                                if len(values) == 1:
                                    values = value.split('2н')
                                #если пара с номером недели в расписании не соответсвует номеру нужной недели, изменения не вносятся, ячейка пропускается
                                if (teacher not in values[0] and week_num == '1') or (teacher not in values[1] and week_num == '2'):
                                    continue
                            
                            if curr_cell.value == None:
                                curr_cell.value = f'{group}\n'
                            else:
                                curr_cell.value = curr_cell.value+f'{group}\n'
            if i+1 < 10:
                curr_date = f'0{i+1}.{date[0]}.{date[1]}'
            else:
                curr_date = f'{i+1}.{date[0]}.{date[1]}'
            
            if curr_date in dates: 
                for row_index in range(dates[curr_date][0], dates[curr_date][1]):
                    if teacher in str(ws_ch._get_cell(row=row_index, column=5).value).split() and ws_ch._get_cell(row=row_index, column=5).value != ws_ch._get_cell(row=row_index, column=7).value:
                        group_curr = ws_ch._get_cell(row=row_index, column=3).value.split()
                        group_curr = f'{group_curr[0]} - {group_curr[-1]}'
                        lesson_num = ws_ch._get_cell(row=row_index, column=2).value
                        group_list = ws_base._get_cell(row=curr_str+int(lesson_num)-1, column=teacher_num+3).value
                        
                        if group_list != None:
                            group_list = group_list.split('\n')
                        else:
                            group_list = ''

                        if group_curr in group_list:
                            group_list.remove(group_curr)
                            text = '\n'.join(group for group in group_list if group != '')
                            ws_base._get_cell(row=curr_str+int(lesson_num)-1, column=teacher_num+3).value = f'{text}\n'
                            

                    elif teacher in str(ws_ch._get_cell(row=row_index, column=7).value).split() and ws_ch._get_cell(row=row_index, column=5).value != ws_ch._get_cell(row=row_index, column=7).value:
                        group_curr = ws_ch._get_cell(row=row_index, column=3).value.split()
                        group_curr = f'{group_curr[0]} - {group_curr[-1]}'
                        if '-' not in str(ws_ch._get_cell(row=row_index, column=2).value):
                            lesson_num = ws_ch._get_cell(row=row_index, column=2).value
                            group_list = ws_base._get_cell(row=curr_str+int(lesson_num)-1, column=teacher_num+3).value
                        else:
                            #!!!!!!!!!!!!!!!!!! -> Номер пары f в формате X-Y
                            continue    

                        if group_list != None:
                            ws_base._get_cell(row=curr_str+int(lesson_num)-1, column=teacher_num+3).value = group_list+f'{group_curr}\n'
                        else:
                            ws_base._get_cell(row=curr_str+int(lesson_num)-1, column=teacher_num+3).value = f'{group_curr}\n'

        if day_num == 6:
            day_num = 0
        else:
            day_num += 1

        curr_str += 6
    
    else:
        wb_base.save(filename='tabel.xlsx')
        wb_base.close()
        wb_gen.close()


def get_start_day(document_path: str) -> tuple:
    doc = Document(document_path)

    for num in paragraphs:
        text = doc.paragraphs[num].text.split()
        date = text[0]

        if '01' == date.split('.')[0]:
            days = monthrange(year=int(date.split('.')[-1]), month=int(date.split('.')[1]))
            return days[1], text[1], date, days[0]
    else:
        for num in paragraphs:
            text = doc.paragraphs[num].text.split()
            date = text[0]

            if '02' == date.split('.')[0]:
                days = monthrange(year=int(date.split('.')[-1]), month=int(date.split('.')[1]))
                return days[1], text[1], date, days[0]


def get_day(date: str, week_num: str) -> str:
    day = week_days[weekday(year=int(date.split('.')[-1]), month=int(date.split('.')[1]), day=int(date.split('.')[0]))]
    return f'{date}\n({day} {week_num}н)'

#удаление лишних пробелов из строки
def del_spaces(string: str) -> str:
    if string != None:
        return ' '.join([str(i) for i in string.split()])

#удаление всех пробелов из строки
def del_all_spaces(string: str) -> str:
    if string != None:
        return string.replace(' ', '')

#удаление лишних символов из названия группы (из общей таблицы с расписанием)
def get_group(group_name: str) -> str:
    group = group_name.split('(')[-1]
    return group[:-1]

def get_group_name(group_name:str) -> str:
    if '-' not in group_name:
        return f'{group_name.split()[0]} - {group_name.split()[1]}'
    else:
        return group_name


def get_doc_tables_dict() -> dict:
    files_amount = len(glob(r'resources\docs\*.docx'))

    for doc_num in range(files_amount):
        doc = Document(f'resources\docs\{doc_num+1}.docx')
        tables = doc.tables

        if len(tables) >= 6:
            tables = tables[:-1]

        for table_num in range(len(tables)):
            table = tables[table_num]
            date = doc.paragraphs[paragraphs[table_num]].text.split()[0]
            doc_tables[date] = table

def docx_to_excel() -> None:
    wb = Workbook()
    ws = wb.active

    ex_row = 0
    ex_col = 1

    for date in doc_tables:
        table = doc_tables[date]
        start_ex_row = ex_row
        for doc_row in table.rows:
            ex_row += 1
            for doc_cell in doc_row.cells:
                ex_col += 1
                text = str(doc_cell.text)
                ws._get_cell(row=ex_row, column=1).value = str(date)

                if '=' not in text and text != '7' and text != '8':
                    if '.' in text:
                        text = text.replace('.', '. ')
                    ws._get_cell(row=ex_row, column=ex_col).value = str(text)
                else:
                    ws._get_cell(row=ex_row, column=ex_col).value = '-'
            ex_col = 1
        dates[date] = [start_ex_row+1, ex_row+1]
    else: 
        wb.save(r'resources\tables\changes.xlsx')
        wb.close() 

def main():
    s = time.time()
    get_doc_tables_dict()
    docx_to_excel()
    script(gen_excel_table_path=r'resources\tables\RASPISANIE.xlsx', document_path=r'resources\docs\1.docx')
    print(time.time()-s)




main()